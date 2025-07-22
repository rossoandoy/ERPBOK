"""
Document processing service for ERPFTS Phase1 MVP

Handles document ingestion, processing, and metadata extraction
with performance optimization through caching, rate limiting, and monitoring.
Supports PDF, DOCX, HTML, and TXT files.
"""

import hashlib
import mimetypes
from pathlib import Path
from typing import Dict, Any, List, Optional, BinaryIO
from uuid import uuid4

from loguru import logger
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from sqlalchemy.orm import Session

from ..core.config import settings
from ..core.exceptions import DocumentProcessingError, ValidationError, RateLimitExceeded
from ..core.cache import get_cache_manager
from ..core.rate_limiter import get_rate_limiter
from ..core.performance import measure_performance, get_performance_monitor
from ..db.session import get_db_session
from ..models.database import Document, KnowledgeChunk
from ..schemas.documents import DocumentCreate, ChunkCreate
from ..utils.text_processing import clean_text, chunk_text, detect_language, calculate_text_hash
from ..utils.file_utils import save_uploaded_file, get_file_info


class DocumentService:
    """Service for document processing and management with performance optimization."""
    
    def __init__(self, db: Optional[Session] = None):
        """Initialize document service with performance components."""
        self.db = db or next(get_db_session())
        self.cache_manager = get_cache_manager()
        self.rate_limiter = get_rate_limiter()
        self.performance_monitor = get_performance_monitor()
        
    async def process_document(
        self, 
        file: BinaryIO, 
        filename: str, 
        user_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Document:
        """
        Process an uploaded document file with rate limiting and performance monitoring.
        
        Args:
            file: File object to process
            filename: Original filename
            user_id: ID of user uploading the document
            metadata: Additional metadata
            
        Returns:
            Created Document object
            
        Raises:
            DocumentProcessingError: If processing fails
            ValidationError: If file validation fails
            RateLimitExceeded: If upload rate limit is exceeded
        """
        async with measure_performance("document_service.process_document", {"filename": filename}):
            try:
                # Rate limiting check for uploads
                if user_id:
                    is_allowed, rate_info = await self.rate_limiter.check_upload_limit(user_id)
                    if not is_allowed:
                        raise RateLimitExceeded(
                            "Document upload rate limit exceeded",
                            retry_after=rate_info.get("retry_after", 60)
                        )
                
                # Validate file
                file_info = await self._validate_file(file, filename)
                
                # Save uploaded file
                file_path = save_uploaded_file(file, filename)
                
                # Extract text content
                text_content = await self._extract_text(file_path, file_info["content_type"])
                
                # Create document record
                document_data = DocumentCreate(
                    filename=filename,
                    file_path=str(file_path),
                    content_type=file_info["content_type"],
                    file_size=file_info["size"],
                    content_hash=calculate_text_hash(text_content),
                    processing_status="processing",
                    source_type="upload",
                    language=detect_language(text_content),
                    metadata=metadata or {},
                    uploaded_by=user_id
                )
                
                document = Document(**document_data.dict(exclude_unset=True))
                document.id = str(uuid4())
                
                self.db.add(document)
                self.db.flush()  # Get the document ID
                
                # Process text into chunks
                await self._create_chunks(document, text_content)
                
                # Update status to completed
                document.processing_status = "completed"
                self.db.commit()
                
                logger.info(f"Successfully processed document: {filename}")
                return document
            
        except Exception as e:
            self.db.rollback()
            logger.error(f"Failed to process document {filename}: {str(e)}")
            raise DocumentProcessingError(f"Document processing failed: {str(e)}")
    
    async def _validate_file(self, file: BinaryIO, filename: str) -> Dict[str, Any]:
        """
        Validate uploaded file.
        
        Args:
            file: File object
            filename: Original filename
            
        Returns:
            File information dictionary
            
        Raises:
            ValidationError: If validation fails
        """
        # Check file size
        file.seek(0, 2)  # Seek to end
        size = file.tell()
        file.seek(0)  # Reset to beginning
        
        max_size = settings.max_file_size_mb * 1024 * 1024
        if size > max_size:
            raise ValidationError(f"File size {size} exceeds limit {max_size}")
        
        # Check file type
        content_type, _ = mimetypes.guess_type(filename)
        if not content_type:
            # Try to detect from file content
            file_header = file.read(1024)
            file.seek(0)
            
            if file_header.startswith(b'%PDF'):
                content_type = 'application/pdf'
            elif b'PK\x03\x04' in file_header[:4]:
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                content_type = 'text/plain'
        
        allowed_types = [
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain',
            'text/html',
            'application/msword'
        ]
        
        if content_type not in allowed_types:
            raise ValidationError(f"File type {content_type} not supported")
        
        return {
            "size": size,
            "content_type": content_type,
            "filename": filename
        }
    
    async def _extract_text(self, file_path: Path, content_type: str) -> str:
        """
        Extract text from file based on content type.
        
        Args:
            file_path: Path to the file
            content_type: MIME type of the file
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentProcessingError: If text extraction fails
        """
        try:
            if content_type == 'application/pdf':
                return await self._extract_pdf_text(file_path)
            elif 'word' in content_type or content_type.endswith('document'):
                return await self._extract_docx_text(file_path)
            elif content_type == 'text/html':
                return await self._extract_html_text(file_path)
            else:  # text/plain and others
                return await self._extract_plain_text(file_path)
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            raise DocumentProcessingError(f"Text extraction failed: {str(e)}")
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                text_parts = []
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                return '\n'.join(text_parts)
                
        except Exception as e:
            logger.error(f"PDF text extraction failed: {str(e)}")
            raise DocumentProcessingError(f"PDF text extraction failed: {str(e)}")
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {str(e)}")
            raise DocumentProcessingError(f"DOCX text extraction failed: {str(e)}")
    
    async def _extract_html_text(self, file_path: Path) -> str:
        """Extract text from HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Get text content
                text = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                return text
                
        except Exception as e:
            logger.error(f"HTML text extraction failed: {str(e)}")
            raise DocumentProcessingError(f"HTML text extraction failed: {str(e)}")
    
    async def _extract_plain_text(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
                
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise DocumentProcessingError("Could not decode text file")
            
        except Exception as e:
            logger.error(f"Plain text extraction failed: {str(e)}")
            raise DocumentProcessingError(f"Plain text extraction failed: {str(e)}")
    
    async def _create_chunks(self, document: Document, text_content: str) -> None:
        """
        Create text chunks from document content.
        
        Args:
            document: Document object
            text_content: Full text content
        """
        try:
            # Clean the text
            cleaned_text = clean_text(text_content)
            
            # Split into chunks
            chunks = chunk_text(
                cleaned_text,
                chunk_size=settings.chunk_size,
                chunk_overlap=settings.chunk_overlap
            )
            
            # Create chunk records
            for idx, (chunk_text, start_pos) in enumerate(chunks):
                chunk = KnowledgeChunk(
                    id=str(uuid4()),
                    document_id=document.id,
                    chunk_index=idx,
                    content=chunk_text,
                    content_hash=calculate_text_hash(chunk_text),
                    start_position=start_pos,
                    end_position=start_pos + len(chunk_text),
                    token_count=len(chunk_text.split()),  # Rough estimate
                    language=document.language
                )
                
                self.db.add(chunk)
            
            logger.info(f"Created {len(chunks)} chunks for document {document.filename}")
            
        except Exception as e:
            logger.error(f"Chunk creation failed: {str(e)}")
            raise DocumentProcessingError(f"Chunk creation failed: {str(e)}")
    
    def get_document(self, document_id: str) -> Optional[Document]:
        """
        Get document by ID.
        
        Args:
            document_id: Document ID
            
        Returns:
            Document object or None
        """
        return self.db.query(Document).filter(Document.id == document_id).first()
    
    def list_documents(
        self, 
        user_id: Optional[str] = None,
        status: Optional[str] = None,
        limit: int = 50,
        offset: int = 0
    ) -> List[Document]:
        """
        List documents with filtering.
        
        Args:
            user_id: Filter by user ID
            status: Filter by processing status
            limit: Maximum number of results
            offset: Offset for pagination
            
        Returns:
            List of Document objects
        """
        query = self.db.query(Document)
        
        if user_id:
            query = query.filter(Document.uploaded_by == user_id)
        
        if status:
            query = query.filter(Document.processing_status == status)
        
        return query.offset(offset).limit(limit).all()
    
    def delete_document(self, document_id: str) -> bool:
        """
        Delete document and associated chunks.
        
        Args:
            document_id: Document ID
            
        Returns:
            True if deleted, False if not found
        """
        try:
            document = self.get_document(document_id)
            if not document:
                return False
            
            # Delete associated chunks
            self.db.query(KnowledgeChunk).filter(
                KnowledgeChunk.document_id == document_id
            ).delete()
            
            # Delete document
            self.db.delete(document)
            self.db.commit()
            
            logger.info(f"Deleted document: {document_id}")
            return True
            
        except Exception as e:
            self.db.rollback()
            logger.error(f"Failed to delete document {document_id}: {str(e)}")
            return False